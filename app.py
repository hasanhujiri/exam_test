import re
import io
import json, base64
from dataclasses import dataclass
from typing import List, Dict

import streamlit as st

# ---- Extractors ----
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

try:
    from docx import Document  # python-docx
except Exception:
    Document = None


@dataclass
class MCQ:
    number: int
    question: str
    options: Dict[str, str]
    answer: str
    explanation: str


# -------- Parsing Logic --------
# Lookaheads ensure every field stops at the NEXT marker (even with no newline).
# Accept both ASCII ':' and full-width '：'.
QUESTION_BLOCK = re.compile(
    r"""
    QUESTION\s*[:：]?\s*(?P<num>\d+)\s*(?P<q>.*?)                     # number + question

    Option\s*A\s*[:：]\s*(?P<A>.*?)(?=\s*Option\s*B\s*[:：])          # stop at 'Option B'
    Option\s*B\s*[:：]\s*(?P<B>.*?)(?=\s*Option\s*C\s*[:：])          # stop at 'Option C'
    Option\s*C\s*[:：]\s*(?P<C>.*?)(?=\s*Option\s*D\s*[:：])          # stop at 'Option D'
    Option\s*D\s*[:：]\s*(?P<D>.*?)(?=\s*Correct\s*Answer\s*[:：])    # stop at 'Correct Answer'

    Correct\s*Answer\s*[:：]\s*(?P<ans>[ABCD])\s*                     # single letter

    Explanation/Reference\s*[:：]\s*(?P<exp>.*?)(?=\s*(?:Q(?:UESTION)?\s*[:：]?\s*\d+)|\Z)
    """,
    re.DOTALL | re.VERBOSE | re.IGNORECASE,
)

CLEAN_CORRECT_IN_OPTION = re.compile(r"\s*Correct\s*Answer\s*[:：]\s*[ABCD]\s*$", re.IGNORECASE | re.DOTALL)
URL_TRAILER = re.compile(r"https?://\S+\s*$", re.IGNORECASE)


def _clean(s: str) -> str:
    return s.replace("\r", "").strip()


# --- Pre-normalization to fix glued markers and page breaks ---
def normalize_text(txt: str) -> str:
    t = txt.replace("\x0c", "\n")  # form feed -> newline
    # kill "Select one:" prompts which can interrupt parsing
    t = re.sub(r"\bSelect\s*one\s*:\s*", "", t, flags=re.IGNORECASE)
    # Ensure markers start on their own line (insert newline if missing)
    markers = [
        r"QUESTION\s*[:：]?\s*\d+",
        r"Option\s*A\s*[:：]",
        r"Option\s*B\s*[:：]",
        r"Option\s*C\s*[:：]",
        r"Option\s*D\s*[:：]",
        r"Correct\s*Answer\s*[:：]\s*[ABCD]",
        r"Explanation/Reference\s*[:：]",
    ]
    for pat in markers:
        t = re.sub(rf"(?<!\n)({pat})", r"\n\1", t, flags=re.IGNORECASE)

    # collapse excessive spaces but keep line structure
    t = re.sub(r"[ \t]+", " ", t)

    # remove URL tails that appear on their own line or glued to the end of options/explanations
    t = re.sub(r"\nhttps?://\S+\s*(?=\n)", "\n", t, flags=re.IGNORECASE)
    t = re.sub(r"https?://\S+\s*$", "", t, flags=re.IGNORECASE)

    return t


def parse_questions_from_text(txt: str) -> List[MCQ]:
    txt = normalize_text(txt)
    results: List[MCQ] = []
    for m in QUESTION_BLOCK.finditer(txt):
        num = int(m.group("num"))
        q = _clean(m.group("q"))
        options = {letter: _clean(m.group(letter)) for letter in ["A", "B", "C", "D"]}

        # Safety cleanup: strip stray "Correct Answer:" fragments or URL tails accidentally glued to options
        for k in ("A", "B", "C", "D"):
            options[k] = CLEAN_CORRECT_IN_OPTION.sub("", options[k]).strip()
            options[k] = URL_TRAILER.sub("", options[k]).strip()

        ans = m.group("ans").strip().upper()
        exp = _clean(m.group("exp"))
        results.append(MCQ(num, q, options, ans, exp))
    results.sort(key=lambda x: x.number)
    return results


# -------- Input Readers ---------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    if pdf_extract_text is None:
        raise RuntimeError("pdfminer.six is required. Install with: pip install pdfminer.six")
    return pdf_extract_text(io.BytesIO(file_bytes))


def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    doc = Document(io.BytesIO(file_bytes))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def parse_any(file_bytes: bytes, filename: str) -> List[MCQ]:
    name = filename.lower()
    if name.endswith(".pdf"):
        txt = extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        txt = extract_text_from_docx(file_bytes)
    else:
        raise RuntimeError("Unsupported file type. Please upload a PDF or DOCX file.")
    return parse_questions_from_text(txt)


# --------- UI ---------
st.set_page_config(page_title="PDF/DOCX ➜ Exam", layout="centered")
st.title("📄 ➜ 📝 Build an Exam from a PDF/DOCX")
st.caption("Upload your Questions & Answers file (PDF or Word) and take an interactive exam.")

with st.sidebar:
    shuffle = st.toggle("Shuffle questions", value=False)
    show_explanations = st.toggle("Show explanations after submit", value=True)

    # One-click cache clear in case the regex/normalizer changes
    if st.button("Clear parser cache"):
        try:
            load_questions.clear()  # type: ignore
            st.success("Cache cleared. Re-upload the file.")
        except Exception:
            st.info("Parser cache not initialized yet.")


@st.cache_data(show_spinner="Extracting text and parsing questions...")
def load_questions(file_bytes: bytes, filename: str):
    return parse_any(file_bytes, filename)


uploaded = st.file_uploader("Upload PDF or Word (.docx)", type=["pdf", "docx"])

if uploaded is not None:
    bytes_data = uploaded.getvalue()
    file_id = f"{uploaded.name}_{len(bytes_data)}"

    try:
        base_questions = load_questions(bytes_data, uploaded.name)
    except Exception as e:
        st.error(f"Parsing failed: {e}")
        st.stop()

    if not base_questions:
        st.warning("No questions were parsed. Please check the file formatting or adjust regex.")
        st.stop()

    # -----------------
    # Stable ordering across reruns
    # -----------------
    if "quiz" not in st.session_state or st.session_state.quiz.get("file_id") != file_id:
        st.session_state.quiz = {
            "file_id": file_id,
            "answers": {},        # num -> 'A'|'B'|'C'|'D'
            "index": 0,
            "submitted": False,
            "order": None,
            "order_mode": None,
        }

    qstate = st.session_state.quiz

    n = len(base_questions)
    if shuffle:
        if not qstate.get("order") or len(qstate["order"]) != n or qstate.get("order_mode") != "shuffled":
            import random
            order = list(range(n))
            rng = random.Random(file_id)  # deterministic per file
            rng.shuffle(order)
            qstate["order"] = order
            qstate["order_mode"] = "shuffled"
    else:
        if not qstate.get("order") or qstate.get("order_mode") != "sequential" or len(qstate["order"]) != n:
            qstate["order"] = list(range(n))
            qstate["order_mode"] = "sequential"

    questions = [base_questions[i] for i in qstate["order"]]

    # -----------------
    # Deep-link / autosave helpers (URL query params)
    # -----------------
    def _encode_answers_map(ans: dict) -> str:
        try:
            return base64.urlsafe_b64encode(json.dumps(ans, ensure_ascii=False).encode("utf-8")).decode("ascii")
        except Exception:
            return ""

    def _decode_answers_map(s: str) -> dict:
        try:
            raw = base64.urlsafe_b64decode(s.encode("ascii")).decode("utf-8")
            d = json.loads(raw)
            return {int(k): v for k, v in d.items()}
        except Exception:
            return {}

    # Restore from URL if it matches
    qp = st.query_params
    if qp.get("file") == file_id:
        try:
            qstate["index"] = int(qp.get("idx", qstate.get("index", 0)))
        except Exception:
            pass
        if "ans" in qp and not qstate["answers"]:
            restored = _decode_answers_map(qp.get("ans", ""))
            if restored:
                qstate["answers"].update(restored)

    def _persist_to_url():
        st.query_params.update({
            "file": file_id,
            "idx": str(qstate.get("index", 0)),
            "ans": _encode_answers_map(qstate.get("answers", {})),
        })

    # -----------------
    # Rendering
    # -----------------
    total = len(questions)
    idx = int(qstate.get("index", 0))
    idx = max(0, min(idx, total - 1))

    # Progress bar + quick stats
    attempted = sum(1 for q in questions if qstate["answers"].get(q.number))
    st.progress(int(100 * attempted / max(1, total)))
    st.caption(f"Progress: {attempted}/{total} answered")

    # Manual backup/restore (optional)
    with st.expander("💾 Backup/Restore progress", expanded=False):
        colB, colL = st.columns(2)
        with colB:
            backup = {
                "file": file_id,
                "index": qstate.get("index", 0),
                "answers": qstate.get("answers", {}),
                "submitted": qstate.get("submitted", False),
                "order": qstate.get("order", []),
                "order_mode": qstate.get("order_mode", "sequential"),
            }
            st.download_button("Download progress JSON",
                               data=json.dumps(backup, ensure_ascii=False, indent=2),
                               file_name="exam_progress.json")
        with colL:
            up = st.file_uploader("Load progress JSON", type=["json"], key="progress_loader")
            if up is not None:
                try:
                    data = json.loads(up.getvalue().decode("utf-8"))
                    if data.get("file") == file_id:
                        qstate["index"] = int(data.get("index", 0))
                        qstate["answers"] = {int(k): v for k, v in data.get("answers", {}).items()}
                        qstate["submitted"] = bool(data.get("submitted", False))
                        _persist_to_url()
                        st.success("Progress loaded.")
                        st.rerun()
                    else:
                        st.warning("Progress file is for a different upload.")
                except Exception as e:
                    st.error(f"Failed to load progress: {e}")

    # Navigator grid (inside an expander)
    with st.expander("🧭 Question navigator", expanded=False):
        per_row = 20
        rows = (total + per_row - 1) // per_row
        for r in range(rows):
            cols = st.columns(per_row)
            for c in range(per_row):
                i = r * per_row + c
                if i >= total:
                    break
                qnum = questions[i].number
                ans_letter = qstate["answers"].get(qnum)
                if ans_letter:
                    if qstate.get("submitted"):
                        badge = "✅" if ans_letter == questions[i].answer else "❌"
                    else:
                        badge = "•"
                else:
                    badge = ""
                label = f"{i+1}{badge}"
                if cols[c].button(label, key=f"jump_{i}"):
                    qstate["index"] = i
                    _persist_to_url()
                    st.rerun()

    st.markdown(f"**Question {idx+1} / {total}**")
    mcq = questions[idx]

    def render_question(mcq: MCQ):
        st.markdown(f"### Question {mcq.number}")
        st.write(mcq.question)

        # Options
        opts = [(l, mcq.options.get(l, "").strip()) for l in ["A", "B", "C", "D"] if mcq.options.get(l, "").strip()]
        if not opts:
            st.warning("⚠️ No options found for this question.")
            return
        labels = [f"{l}. {t}" for l, t in opts]

        prev_letter = qstate["answers"].get(mcq.number)
        try:
            default_index = [l for l, _ in opts].index(prev_letter) if prev_letter else 0
        except ValueError:
            default_index = 0

        choice_label = st.radio(
            "Select one:",
            labels,
            index=default_index,
            key=f"radio_{mcq.number}",
        )
        if choice_label:
            qstate["answers"][mcq.number] = choice_label[0]
            _persist_to_url()

        # Correctness after submit
        if qstate.get("submitted"):
            user = qstate["answers"].get(mcq.number)
            correct = mcq.answer
            if user == correct:
                st.success(f"✅ Correct: {correct}")
            else:
                st.error(f"❌ Your answer: {user} | Correct: {correct}")

        # Reveal control
        with st.expander("⬇️ Show Answer & Explanation"):
            correct = mcq.answer
            opt_text = mcq.options.get(correct, "")
            st.markdown(f"**Correct Answer:** {correct}. {opt_text}")
            if mcq.explanation and (show_explanations or qstate.get("submitted")):
                st.markdown("**Explanation:**")
                st.write(mcq.explanation)

    render_question(mcq)

    # Navigation (only place index changes)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("⟵ Previous", use_container_width=True, disabled=idx == 0):
            qstate["index"] = max(0, idx - 1)
            _persist_to_url()
            st.rerun()
    with c2:
        if st.button("Next ⟶", use_container_width=True, disabled=idx >= total - 1):
            qstate["index"] = min(total - 1, idx + 1)
            _persist_to_url()
            st.rerun()
    with c3:
        if st.button("Submit", type="primary", use_container_width=True, disabled=qstate.get("submitted", False)):
            qstate["submitted"] = True
            _persist_to_url()
            st.rerun()
    with c4:
        if st.button("Restart", use_container_width=True):
            st.session_state.quiz = {
                "file_id": file_id,
                "answers": {},
                "index": 0,
                "submitted": False,
                "order": qstate.get("order"),
                "order_mode": qstate.get("order_mode"),
            }
            _persist_to_url()
            st.rerun()

    # Score + export after submit
    if qstate.get("submitted"):
        correct = 0
        attempted = 0
        for item in questions:
            user = qstate["answers"].get(item.number)
            if user:
                attempted += 1
                if user == item.answer:
                    correct += 1
        st.info(f"Score: {correct}/{total} | Attempted: {attempted}")

        out = [q.__dict__ for q in questions]
        st.download_button(
            "Download parsed questions as JSON",
            data=json.dumps(out, ensure_ascii=False, indent=2).encode("utf-8"),
            file_name="parsed_questions.json",
            mime="application/json",
        )
else:
    st.info("Upload a PDF or DOCX to get started.")

